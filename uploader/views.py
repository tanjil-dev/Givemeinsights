import matplotlib, os, io
matplotlib.use('Agg')
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.core.files.storage import FileSystemStorage
from itertools import combinations
from django.shortcuts import render
from ydata_profiling import ProfileReport
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report, confusion_matrix
import seaborn as sns
import docx, itertools
from bs4 import BeautifulSoup
from wordcloud import WordCloud
import pandas as pd
import matplotlib.pyplot as plt
import io, spacy, re, unicodedata
import numpy as np
import base64, string
from collections import Counter
from django.views.generic import TemplateView
from datetime import datetime
from io import StringIO, BytesIO
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from django.conf import settings
from .forms import *
from django.views import View
from django.core.mail import send_mail
import logging
import textstat
import docx2txt
from textblob import TextBlob
from django.http import HttpResponse
from tempfile import NamedTemporaryFile
import plotly.express as px


logger = logging.getLogger(__name__)

class WordAnalysisView(TemplateView):
    template_name = 'word-analysis.html'
class ExcelAnalysisView(TemplateView):
    template_name = 'excel-analysis.html'

class HowItWorks(TemplateView):
    template_name = 'how_it_works.html'
@method_decorator(csrf_exempt, name='dispatch')
class ContactView(View):
    def get(self, request):
        # Display the empty contact form
        form = ContactForm()
        return render(request, 'contact.html', {'form': form})

    def post(self, request):
        form = ContactForm(request.POST)
        if form.is_valid():
            name = form.cleaned_data['name']
            email = form.cleaned_data['email']
            phone = form.cleaned_data['phone']
            subject = form.cleaned_data['subject']
            message = form.cleaned_data['message']

            try:
                # Send the email to the client's Gmail
                send_mail(
                    subject=f"New message from {name}",
                    message=f"Name: {name}\nEmail: {email}\nPhone: {phone}\nSubject: {subject}\nMessage:\n{message}",
                    from_email=email,
                    recipient_list=[settings.CLIENT_EMAIL],
                    fail_silently=False
                )

                # Return the success message
                success_message = 'Thank you for your message! We will get back to you shortly.'
                return render(request, 'contact.html', {'form': ContactForm(), 'success_message': success_message})
            except Exception as e:
                logger.error(f"Error sending email: {e}")
                return HttpResponse("There was an error sending your message. Please try again later.", status=500)

        return render(request, 'contact.html', {'form': form})

#list of common stopwords
stop_words = set([
    "i", "me", "my", "myself", "we", "our", "ours", "ourselves", "you", "your", "yours", 
    "yourself", "yourselves", "he", "him", "his", "himself", "she", "her", "hers", "herself", 
    "it", "its", "itself", "they", "them", "their", "theirs", "themselves", "what", "which", 
    "who", "whom", "this", "that", "these", "those", "am", "is", "are", "was", "were", 
    "be", "been", "being", "have", "has", "had", "having", "do", "does", "did", "doing", "a", 
    "an", "the", "and", "but", "if", "or", "because", "as", "until", "while", "of", "at", 
    "by", "for", "with", "about", "against", "between", "into", "through", "during", 
    "before", "after", "above", "below", "to", "from", "up", "down", "in", "out", 
    "on", "off", "over", "under", "again", "further", "then", "once", "here", "there", 
    "when", "where", "why", "how", "all", "any", "both", "each", "few", "more", "most", 
    "other", "some", "such", "no", "nor", "not", "only", "own", "same", "so", "than", 
    "too", "very", "s", "t", "can", "will", "just", "don", "should", "now", "d", "ll", 
    "m", "o", "re", "ve", "y", "ain", "aren", "couldn", "didn", "doesn", "hadn", 
    "hasn", "haven", "isn", "ma", "mightn", "mustn", "needn", "shan", "shouldn", 
    "wasn", "weren", "won", "wouldn", "like"
])

nlp = spacy.load("en_core_web_sm")

def extract_text(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

@csrf_exempt
def phrases_used_view(request):
    form = DocumentUploadForm()
    grouped_phrases_by_count = []

    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            # Extract text from the uploaded document
            text = extract_text(request.FILES['file'])
            word_list = text.split()
            grouped_words = []

            # Generate phrases of increasing sizes based on the document's content
            max_phrase_size = 15  # This ensures phrases of any length from 1 to the document's length
            for group_size in range(1, max_phrase_size + 1):  # Dynamic loop from 1-word phrases to max_phrase_size
                for k in range(len(word_list) - group_size + 1):
                    group_slice = word_list[k:k + group_size]
                    phrase = " ".join(group_slice)
                    grouped_words.append(phrase)

            # Convert to DataFrame for easier manipulation
            df = pd.DataFrame(grouped_words, columns=['phrase'])

            # Count the occurrences of each phrase
            phrase_counts = df['phrase'].value_counts().reset_index()
            phrase_counts.columns = ['phrase', 'count']  # Rename columns explicitly
            phrase_counts = phrase_counts[phrase_counts['count'] > 1]  # Only keep phrases that appear more than once

            # Add a column for the number of words in each phrase
            phrase_counts['number_of_words'] = phrase_counts['phrase'].apply(lambda x: len(x.split()))

            # Dynamically calculate the maximum word count in the document
            max_word_count = phrase_counts['number_of_words'].max()  # Max word count in phrases

            # Group and limit top 15 for each word count dynamically
            for i in range(1, max_word_count + 1):  # From 1 word to the longest phrase length
                subset = phrase_counts[phrase_counts['number_of_words'] == i].copy()
                if not subset.empty:
                    grouped_phrases_by_count.append({
                        'word_count': i,
                        'phrases': subset.head(15).to_dict('records')  # Adjust the top 15 as needed
                    })

    return render(request, 'phrases_used.html', {
        'form': form,
        'grouped_phrases_by_count': grouped_phrases_by_count
    })


# Normalize special characters and quotes
def normalize_text(text):
    text = unicodedata.normalize("NFKD", text)
    text = text.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'")
    text = text.replace("′", "'").replace("″", '"')  # Prime and double-prime to ASCII
    return text.strip()

def is_valid_entity(text):
    text = normalize_text(text)

    # Reject very short or non-alphabetic
    if len(text) <= 2 or not any(c.isalpha() for c in text):
        return False

    # Clean punctuation for blacklist checking
    clean_text = re.sub(r"[^\w\s]", "", text.lower())  # remove punctuation
    banned = {"nt", "n", "itt", "i", "s", "ll", "ve"}

    if clean_text in banned:
        return False

    # Filter coordinates or broken patterns
    if re.search(r"\d+[a-zA-Z]*[\s′″\"']+", text):  # like 62o 17′ 20″
        return False

    # Filter initials or single-letter followed by punctuation or quote
    if re.match(r"^[A-Z]\.\s*[\"']?$", text):
        return False

    return True

@csrf_exempt
def labels_view(request):
    form = DocumentUploadForm()
    grouped_labels = {}

    if request.method == 'POST':
        form = DocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            text = extract_text(request.FILES['file'])
            nlp.max_length = 2030000  # set max length if needed for large docs
            doc = nlp(text)

            entity_counter = Counter(
                (ent.label_, ent.text) for ent in doc.ents
                if is_valid_entity(ent.text)
            )

            # Group by entity label
            grouped_labels = {}
            for (label, ent_text), count in entity_counter.items():
                grouped_labels.setdefault(label, []).append((ent_text, count))

            # Sort entities within each group by count
            for label in grouped_labels:
                grouped_labels[label] = sorted(grouped_labels[label], key=lambda x: -x[1])

    return render(request, 'labels.html', {
        'form': form,
        'grouped_labels': grouped_labels
    })


def home(request):
    return render(request, 'home.html')

@csrf_exempt
def upload_excel(request):
    form = UploadFileForm()
    return render(request, 'upload_excel.html', {'form': form})


@csrf_exempt
def upload_docx(request):
    start_time = datetime.now()
    error_message = None
    full_wordcloud_img = None
    top_wordcloud_img = None
    time_taken = None

    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['file']

            if not docx_file.name.endswith('.docx'):
                error_message = "Invalid file type. Please upload a .docx file."
            else:
                doc_text = extract_text_from_docx(docx_file)
                full_wordcloud_img, top_wordcloud_img, time_taken = generate_wordcloud(doc_text, start_time)
    else:
        form = UploadFileForm()
    end_time = datetime.now()
    return render(request, 'word_cloud.html', {
        'form': form,
        'full_wordcloud_img': full_wordcloud_img,
        'top_wordcloud_img': top_wordcloud_img,
        'error_message': error_message,
        'time_taken': time_taken,
        'start_time': start_time,
        'end_time': end_time,
    })

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def generate_wordcloud(text, start_time):
    # Convert text to lowercase and remove punctuation
    text = text.lower()
    translator = str.maketrans('', '', string.punctuation)
    text = text.translate(translator)
    
    # Tokenize the text and remove stopwords manually
    word_tokens = text.split()
    
    # Remove stopwords from the word tokens
    filtered_words = [word for word in word_tokens if word not in stop_words]

    # Count word frequencies
    word_counts = Counter(filtered_words)
    
    # Get the top 10 most common words
    top_10_words = word_counts.most_common(10)
    
    # Prepare the text for the top 10 words word cloud
    top_10_words_text = ' '.join([word[0] for word in top_10_words])
    
    # Generate the full word cloud from filtered words (after removing stopwords)
    full_wordcloud = WordCloud(width=800, height=400, background_color='white', max_words=200).generate(' '.join(filtered_words))
    
    # Generate the top 10 word cloud
    top_wordcloud = WordCloud(width=800, height=400, background_color='white', max_words=10, min_font_size=20).generate(top_10_words_text)
    
    # Save the full word cloud image to a BytesIO object
    full_img_io = io.BytesIO()
    plt.figure(figsize=(10, 5))
    plt.imshow(full_wordcloud, interpolation="bilinear")
    plt.axis('off')
    plt.savefig(full_img_io, format='png')
    full_img_io.seek(0)
    
    # Save the top 10 word cloud image to a BytesIO object
    top_img_io = io.BytesIO()
    plt.figure(figsize=(10, 5))
    plt.imshow(top_wordcloud, interpolation="bilinear")
    plt.axis('off')
    plt.savefig(top_img_io, format='png')
    top_img_io.seek(0)
    
    # Convert images to base64 to embed in HTML
    full_wordcloud_img = base64.b64encode(full_img_io.getvalue()).decode()
    top_wordcloud_img = base64.b64encode(top_img_io.getvalue()).decode()
    
    #calculate total time taken        
    end_time = datetime.now()    
    total_seconds = (end_time - start_time).total_seconds()    
    minutes = total_seconds // 60
    seconds = total_seconds % 60
    time_taken = f"{int(minutes)} minutes and {int(seconds)} seconds"
    return full_wordcloud_img, top_wordcloud_img, time_taken



# Helper function to process the CSV and generate both the boxplot and time series plot
def process_csv_and_generate_plots(csv_file, start_time):
    # Load the CSV file into a DataFrame
    df = pd.read_csv(csv_file, parse_dates=['date'])
    
    # Drop unnecessary columns (store, item)
    df_sales_only = df.drop(['store', 'item'], axis=1)

    # Extract day, month, year, and dayofweek from the 'date' column
    df_sales_only['day'] = df_sales_only['date'].dt.day
    df_sales_only['month'] = df_sales_only['date'].dt.month
    df_sales_only['year'] = df_sales_only['date'].dt.year
    df_sales_only['dayofweek'] = df_sales_only['date'].dt.dayofweek
    day_colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2']

    
    # Boxplot for sales by day of the week
    plt.figure(figsize=(10, 6))
    sns.boxplot(x='dayofweek', y='sales', data=df_sales_only, palette=day_colors)
    plt.title('Sales Distribution by Day of the Week')
    plt.xlabel('Day of the Week')
    plt.ylabel('Sales')
    plt.xticks(ticks=range(7), labels=['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun'])

    # Save the boxplot to a BytesIO object and convert to base64
    img_io = io.BytesIO()
    plt.savefig(img_io, format='png')
    img_io.seek(0)
    boxplot_img = base64.b64encode(img_io.getvalue()).decode()

    # Time Series Plot for sales by store and item (filtering for store == 1 and item == 1)
    store_item_df = df[(df['store'] == 1) & (df['item'] == 1)]
    
    plt.figure(figsize=(10, 6))
    store_item_df.set_index('date')['sales'].plot()
    plt.title('Sales Time Series for Store 1, Item 1')
    plt.xlabel('Date')
    plt.ylabel('Sales')

    # Save the time series plot to a BytesIO object and convert to base64
    ts_img_io = io.BytesIO()
    plt.savefig(ts_img_io, format='png')
    ts_img_io.seek(0)
    ts_img = base64.b64encode(ts_img_io.getvalue()).decode()

    #calculate total time taken        
    end_time = datetime.now()    
    total_seconds = (end_time - start_time).total_seconds()    
    minutes = total_seconds // 60
    seconds = total_seconds % 60
    time_taken = f"{int(minutes)} minutes and {int(seconds)} seconds"
    return boxplot_img, ts_img, time_taken

@csrf_exempt
def upload_csv(request):
    start_time = datetime.now()
    error_message = None
    boxplot_img = None
    ts_img = None
    time_taken = None

    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            csv_file = request.FILES['file']
        
            # Check if the uploaded file is a CSV
            if not csv_file.name.endswith('.csv'):
                error_message = "Invalid file type. Please upload a .csv file."
            else:
                # Process the CSV file and generate both plots
                boxplot_img, ts_img, time_taken = process_csv_and_generate_plots(csv_file=csv_file, start_time=start_time)

    else:
        form = UploadFileForm()
    end_time = datetime.now()
    return render(request, 'upload_csv.html', {
        'form': form,
        'boxplot_img': boxplot_img,
        'ts_img': ts_img,
        'error_message': error_message,
        'time_taken': time_taken,
        'start_time': start_time,
        'end_time': end_time,
    })

@csrf_exempt
def titanic_view(request):
    if request.method == 'POST' and request.FILES['csv_file']:
        # Get the uploaded file
        uploaded_file = request.FILES['csv_file']
        
        # Read the file into a pandas DataFrame without saving to disk
        file_content = uploaded_file.read().decode('utf-8')  # Read the file in memory
        titanic_df = pd.read_csv(StringIO(file_content))  # Read it as CSV directly into a DataFrame

        # Data Processing Logic: Get the data types of the columns
        data_dtypes = titanic_df.dtypes.to_frame(name='Data Type')  # Get the data types of the columns

        # Calculate null values for each column
        null_values = titanic_df.isnull().sum()

        # Create histogram for 'Age' column, excluding null values
        fig, ax = plt.subplots(figsize=(8, 6))
        titanic_df.loc[-titanic_df.Age.isnull(), 'Age'].plot.hist(ax=ax, bins=30, color='#2577B4', edgecolor='black')
        ax.set_title('Histogram of Age (Excluding Null Values)')
        ax.set_xlabel('Age')
        ax.set_ylabel('Frequency')

        # Convert plot to PNG image and encode it in base64 for embedding in HTML
        canvas = FigureCanvas(fig)
        image_stream = io.BytesIO()
        canvas.print_png(image_stream)
        image_stream.seek(0)
        plot_data1 = base64.b64encode(image_stream.read()).decode('utf-8')
        
        # Data Processing Logic: Fill missing 'Age' values with mean Age by 'Sex'
        titanic_df['Age'] = titanic_df['Age'].fillna(titanic_df.groupby('Sex')['Age'].transform('mean'))

        # Create histogram for 'Age' column after filling missing values
        fig, ax = plt.subplots(figsize=(8, 6))
        titanic_df['Age'].hist(ax=ax, bins=30, color='#2577B4', edgecolor='black')
        ax.set_title('Histogram of Age (After Filling Missing Values with Grouped Mean by Sex)')
        ax.set_xlabel('Age')
        ax.set_ylabel('Frequency')

        # Convert plot to PNG image and encode it in base64 for embedding in HTML
        canvas = FigureCanvas(fig)
        image_stream = io.BytesIO()
        canvas.print_png(image_stream)
        image_stream.seek(0)
        plot_data2 = base64.b64encode(image_stream.read()).decode('utf-8')

        # Render to HTML
        return render(request, 'titanic.html', {
            'null_values': null_values.to_frame(name='Null Values').to_html(classes='table table-striped'),
            'data_dtypes': data_dtypes.to_html(classes='table table-striped'),  # Passing dtypes to template
            'plot_data1': plot_data1,
            'plot_data2': plot_data2,
        })
    
    return render(request, 'titanic.html')

@csrf_exempt
def heatmap_view(request):
    eda_html = None
    heatmap_image = None
    histogram_image = None
    data_summary = None
    null_summary = None

    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']

        try:
            # Read uploaded Excel file into DataFrame
            df = pd.read_excel(excel_file)

            # Prepare correlation heatmap
            numeric_df = df.select_dtypes(include=['number'])
            correlation_matrix = numeric_df.corr()

            plt.figure(figsize=(10, 5))
            sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm')
            plt.title("Correlations Table")

            # Save image to memory
            buf = io.BytesIO()
            plt.savefig(buf, format='png', bbox_inches='tight')
            buf.seek(0)
            heatmap_image = base64.b64encode(buf.read()).decode('utf-8')
            buf.close()
            plt.close()

            # === 2. Generate Histogram ===
            numeric_df.hist(figsize=(16, 20), bins=50, xlabelsize=8, ylabelsize=8)
            plt.tight_layout()

            buf2 = io.BytesIO()
            plt.savefig(buf2, format='png', bbox_inches='tight')
            buf2.seek(0)
            histogram_image = base64.b64encode(buf2.read()).decode('utf-8')
            buf2.close()
            plt.close()

            # Optional EDA table (describe numeric columns)
            eda_html = numeric_df.describe().to_html(classes='table table-bordered')

            # Data summary (df.info()) converted to a DataFrame
            # Instead of parsing df.info(), we directly use dtypes
            data_summary = pd.DataFrame({
                'Column': df.columns,
                'Non-Null Count': df.count(),
                'Dtype': df.dtypes
            })

            # Convert to HTML
            data_summary_html = data_summary.to_html(classes='table table-bordered', index=False)
            data_summary = data_summary_html

            # Number of nulls
            null_summary = df.isnull().sum().sort_values(ascending=False).to_frame().reset_index()
            null_summary.columns = ['Titles', 'Null Count']  # Rename the columns

            # Convert to HTML table with custom classes
            null_summary = null_summary.to_html(classes='table table-bordered', index=False)


        except Exception as e:
            eda_html = f"<p class='text-danger'>Error processing file: {e}</p>"

    return render(request, 'heatmap.html', {
        'eda_html': eda_html,
        'heatmap_image': heatmap_image,
        'histogram_image': histogram_image,
        'data_summary': data_summary,
        'null_summary': null_summary
    })



@csrf_exempt
def scatter_plots_view(request):
    scatter_plots = []

    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        df = pd.read_excel(excel_file)

        numeric_df = df.select_dtypes(include=['number'])
        columns = numeric_df.columns
        pairs = list(itertools.combinations(columns, 2))

        for x, y in pairs:
            plt.figure(figsize=(6, 4))
            plt.scatter(numeric_df[x], numeric_df[y], alpha=0.7)
            plt.xlabel(x)
            plt.ylabel(y)
            plt.title(f"{x} vs {y}")

            buffer = BytesIO()
            plt.savefig(buffer, format='png')
            plt.close()
            buffer.seek(0)
            image_png = buffer.getvalue()
            buffer.close()

            image_base64 = base64.b64encode(image_png).decode('utf-8')
            image_uri = f'data:image/png;base64,{image_base64}'
            scatter_plots.append(image_uri)

    return render(request, 'scatter_plots.html', {'scatter_plots': scatter_plots})

def plot_to_base64():
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    buf.seek(0)
    string = base64.b64encode(buf.read())
    return f'data:image/png;base64,{string.decode()}'

@csrf_exempt
def handle_excel_upload(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        df = pd.read_excel(request.FILES['excel_file'])
        numeric_df = df.select_dtypes(include='number')
        return df, numeric_df
    return None, None

@csrf_exempt
def eda_line_graphs(request):
    df, numeric_df = handle_excel_upload(request)
    images = []

    if numeric_df is not None:
        for column in numeric_df.columns:
            plt.figure()
            plt.plot(numeric_df[column], marker='o', linestyle='-')
            plt.title(f'Graph for {column}')
            plt.xlabel('Index')
            plt.ylabel(column)
            plt.grid(True)
            images.append(plot_to_base64())
            plt.close()
    return render(request, 'line_graphs.html', {'line_graphs': images})

@csrf_exempt
def eda_box_plots(request):
    df, numeric_df = handle_excel_upload(request)
    images = []

    if numeric_df is not None:
        for column in numeric_df.columns:
            plt.figure(figsize=(5, 4))
            sns.boxplot(y=numeric_df[column].dropna(), color='skyblue')
            plt.title(f'Boxplot of {column}')
            plt.ylabel('Values')
            plt.grid(True, linestyle='--', alpha=0.3)
            images.append(plot_to_base64())
            plt.close()
    return render(request, 'box_plots.html', {'box_plots': images})

@csrf_exempt
def eda_pair_plot(request):
    df, numeric_df = handle_excel_upload(request)
    image = None

    if numeric_df is not None:
        numeric_df2 = numeric_df.replace([np.inf, -np.inf], np.nan)
        numeric_df2 = numeric_df2.dropna(thresh=2)  # Only drop rows that are almost empty

        if len(numeric_df2.columns) >= 2 and not numeric_df2.empty:
            sns_plot = sns.pairplot(numeric_df2)
            fig = sns_plot.fig
            image = pair_plot_to_base64(fig)
            plt.close(fig)

    return render(request, 'pair_plot.html', {'pair_plot': image})

@csrf_exempt
def pair_plot_to_base64(fig):
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    image_base64 = base64.b64encode(buf.read()).decode('utf-8')
    return f"data:image/png;base64,{image_base64}"
@csrf_exempt
def linear_regression(request):
    regression_plot_urls = []

    if request.method == 'POST' and request.FILES['excel_file']:
        # Handle file upload
        uploaded_file = request.FILES['excel_file']
        fs = FileSystemStorage()
        filename = fs.save(uploaded_file.name, uploaded_file)
        file_url = fs.url(filename)

        # Read the Excel file
        df = pd.read_excel(os.path.join(fs.location, filename))

        # Select numeric columns
        numeric_df = df.select_dtypes(include=[np.number])
        cols = numeric_df.columns

        # Calculate number of rows needed for the 2-column layout
        num_plots = len(list(combinations(cols, 2)))
        num_rows = (num_plots // 2) + (num_plots % 2)

        # Create a grid of 2 columns
        fig, axes = plt.subplots(nrows=num_rows, ncols=2, figsize=(12, 5 * num_rows))
        axes = axes.flatten()  # Flatten axes array for easy iteration

        # Plot each pair of numeric columns
        plot_counter = 0
        for x_col, y_col in combinations(cols, 2):
            x = numeric_df[x_col].values
            y = numeric_df[y_col].values

            # Drop rows with non-finite values
            mask = np.isfinite(x) & np.isfinite(y)
            x = x[mask]
            y = y[mask]

            if len(x) > 1 and len(y) > 1:
                # Perform Linear Regression (y = a * x + b)
                a, b = np.polyfit(x, y, 1)
                y_fit = a * x + b

                # Calculate correlation coefficient (Pearson r)
                corr_matrix = np.corrcoef(x, y)
                corr = corr_matrix[0, 1]

                # Plotting in the corresponding subplot
                ax = axes[plot_counter]
                ax.scatter(x, y, label=f'{y_col} vs {x_col}', color='blue')
                ax.plot(x, y_fit, label=f'Fit: y = {a:.2f}x + {b:.2f}', color='red')
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.set_title(f'{y_col} vs {x_col}\nCorrelation r = {corr:.2f}')
                ax.legend()
                ax.grid(True)

                plot_counter += 1

        # Remove any unused axes (if the number of plots is not a perfect multiple of 2)
        for i in range(plot_counter, len(axes)):
            fig.delaxes(axes[i])

        # Adjust the layout to remove extra space at the top and between plots
        plt.subplots_adjust(hspace=0.3, wspace=0.2, top=0.85, bottom=0.07)  # Further reduce top space
        # Ensure no space at the top
        fig.subplots_adjust(top=0.95, bottom=0.05, left=0.05, right=0.95)  # Adjust top margin directly
        # Save plot to in-memory file and convert to base64
        buffer = BytesIO()
        canvas = FigureCanvas(fig)
        canvas.print_png(buffer)
        plot_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        regression_plot_urls.append(f"data:image/png;base64,{plot_base64}")

        # Don't forget to delete the uploaded file after processing
        os.remove(os.path.join(fs.location, filename))

    return render(request, 'linear_regression.html', {'regression_plots': regression_plot_urls})

class RobotsTxtView(View):
    def get(self, request):
        file_path = os.path.join(settings.BASE_DIR, 'uploader/static', 'robots.txt')
        with open(file_path, 'r') as f:
            return HttpResponse(f.read(), content_type='text/plain')

class SitemapXmlView(View):
    def get(self, request):
        file_path = os.path.join(settings.BASE_DIR, 'uploader/static', 'sitemap.xml')
        with open(file_path, 'r') as f:
            return HttpResponse(f.read(), content_type='application/xml')

@csrf_exempt
def generate_profile_report(request):
    profile_report_html = None
    report_path = None

    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']

        try:
            # Read the uploaded Excel file
            df = pd.read_excel(excel_file)

            # Generate the profile report
            profile = ProfileReport(df, title="My Data Profile Report")

            # Save the report as an HTML file to a temporary location
            report_path = 'profile_report.html'

            # Save the profile to a file on the filesystem
            profile.to_file(report_path)

            # Read the saved report content
            with open(report_path, 'r') as report_file:
                profile_report_html = report_file.read()

            # Use BeautifulSoup to parse and remove the footer content
            soup = BeautifulSoup(profile_report_html, 'html.parser')

            # Remove <footer> tag if it exists
            footer = soup.find('footer')
            if footer:
                footer.decompose()  # Removes the footer element completely

            # Remove <p> tag with "Brought to you by YData" text if it exists
            footer_p_tag = soup.find('p', class_='text-body-secondary text-end')
            if footer_p_tag:
                footer_p_tag.decompose()  # Removes the <p> tag containing YData attribution

            # Also remove any other possible footer-related elements with specific classes or text
            ydata_footer_p = soup.find('p', string=lambda text: text and 'Brought to you by YData' in text)
            if ydata_footer_p:
                ydata_footer_p.decompose()  # Remove this specific YData attribution paragraph

            # Get the updated HTML without the footer or YData attribution
            profile_report_html = str(soup)

        except Exception as e:
            profile_report_html = f"<p class='text-danger'>Error processing file: {e}</p>"

    return render(request, 'profile_report.html',
                  {'profile_report_html': profile_report_html, 'report_path': report_path})


def get_sentiment(text):
    blob = TextBlob(text)
    polarity = blob.sentiment.polarity
    subjectivity = blob.sentiment.subjectivity

    if polarity > 0:
        sentiment = "Positive"
    elif polarity < 0:
        sentiment = "Negative"
    else:
        sentiment = "Neutral"

    # Return sentiment, polarity, and subjectivity
    return pd.Series([sentiment, polarity, subjectivity])

import re

def split_into_sentences(text):
    # This regex splits on period, question mark, or exclamation followed by space or line end
    sentence_endings = re.compile(r'(?<=[.!?])\s+')
    sentences = sentence_endings.split(text)
    sentences = [s.strip() for s in sentences if s.strip()]
    return sentences

def calculate_readability(file_path: str, TopPositiveEntered: int, MostNegative: int, user_threshold: int):
    """
    Reads a .docx file, computes readability + per-sentence sentiment, and
    returns:
        - sentence_df (DataFrame)
        - image_data (Readability Chart as base64 PNG)
        - sentiment_plot_data (Positive vs Negative score histogram, base64 PNG)
        - sentiment_count_plot_data (Countplot of labels, base64 PNG)
        - formatted_top_positive_sentences (list[str])
        - formatted_low_rank_sentences (list[str])   # low readability
        - formatted_top_negative_sentences (list[str])
    """
    # Extract text from .docx
    my_text = docx2txt.process(file_path)

    # Split the text into sentences
    sentences = split_into_sentences(my_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 0]

    if not sentences:
        return None, None, None, None, [], [], []

    # Compute readability per sentence
    readability_scores = [textstat.flesch_reading_ease(sentence) for sentence in sentences]

    # Build DataFrame
    sentence_df = pd.DataFrame({
        'Sentence': sentences,
        'Readability_Score': readability_scores
    })

    # Sentiment analysis
    sentence_df[['Sentiment', 'Sentiment_Score', 'Subjectivity']] = sentence_df['Sentence'].apply(get_sentiment)
    sentence_df['Opinion_Factor'] = sentence_df['Subjectivity']
    sentence_df['word_count'] = sentence_df['Sentence'].apply(lambda x: len(x.split()))
    sentence_df['Sentence_No'] = range(1, len(sentence_df) + 1)

    # Rank sentiment: 1 = most negative, 100 = most positive
    rank_pct = sentence_df['Sentiment_Score'].rank(method='min', ascending=True, pct=True)
    sentence_df['Sentiment_Rank'] = (rank_pct * 99 + 1).round().astype(int)

    # Rank readability: lower score = harder
    sentence_df['Readability_Rank'] = sentence_df['Readability_Score'].rank(
        method='min', ascending=True, pct=True
    )
    sentence_df['Readability_Rank'] = (sentence_df['Readability_Rank'] * 99 + 1).round().astype(int)

    # -----------------------------
    # Low readability sentences
    # -----------------------------
    low_rank_df = sentence_df[sentence_df['Readability_Rank'] < user_threshold]
    if low_rank_df.empty:
        low_rank_df = sentence_df.nsmallest(1, 'Readability_Score')

    formatted_low_rank_sentences = [
        f"Sentence {row['Sentence_No']}: {row['Sentence']}\n"
        f"❌ Not easy to read — Score: {row['Readability_Score']:.2f}, "
        f"Rank: {row['Readability_Rank']}/100"
        for _, row in low_rank_df.iterrows()
    ]

    # -----------------------------
    # Top positive sentences
    # -----------------------------
    TopPositive = 100 - TopPositiveEntered
    filtered_df_positive = sentence_df[sentence_df['Sentiment_Rank'] > TopPositive][
        ['Sentence_No', 'Sentence', 'Sentiment_Score', 'Sentiment_Rank']
    ]
    formatted_top_positive_sentences = [
        f"Sentence {row['Sentence_No']}: {row['Sentence']}\n"
        f"Sentiment Score: {row['Sentiment_Score']:.2f}\n"
        f"🏅 Sentiment Rank: {row['Sentiment_Rank']}/100"
        for _, row in filtered_df_positive.iterrows()
    ]

    # -----------------------------
    # Top negative sentences
    # -----------------------------
    filtered_df_negative = sentence_df[sentence_df['Sentiment_Rank'] < MostNegative][
        ['Sentence_No', 'Sentence', 'Sentiment_Score', 'Sentiment_Rank']
    ]
    formatted_top_negative_sentences = [
        f"Sentence {row['Sentence_No']}: {row['Sentence']}\n"
        f"Sentiment Score: {row['Sentiment_Score']:.2f}\n"
        f"🏅 Sentiment Rank: {row['Sentiment_Rank']}/100"
        for _, row in filtered_df_negative.iterrows()
    ]

    # -----------------------------
    # Plots
    # -----------------------------
    sentiment_plot_data = None
    sentiment_count_plot_data = None

    # -----------------------------
    # Readability Chart (Bar Plot)
    # -----------------------------
    # Update the bins to cover a wider range
    bins = [0, 10, 30, 50, 60, 80, 90, 100]  # More granular bins
    bin_labels = [f"({bins[i]}, {bins[i + 1]})" for i in range(len(bins) - 1)]  # Use bin ranges for labels

    # Bin readability scores into categories
    sentence_df['Readability_Binned'] = pd.cut(sentence_df['Readability_Score'], bins=bins, labels=bin_labels,
                                               right=False)

    # Count the number of sentences in each bin
    bin_counts = sentence_df['Readability_Binned'].value_counts().sort_index()

    # Colors for the bins (More colorful gradient)
    colors = sns.color_palette("RdYlGn", len(bins) - 1)

    # Create the Readability Chart (Bar Plot)
    plt.figure(figsize=(10, 6))
    sns.barplot(x=bin_counts.index.astype(str), y=bin_counts.values, palette=colors)

    # Labels and styling
    plt.xlabel('Readability Range')
    plt.ylabel('Number of Sentences')
    plt.title('Readability of Sentences (Higher is easier to read)')
    plt.xticks(rotation=45)
    plt.grid(axis='y', linestyle='--', alpha=0.6)
    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    image_data = base64.b64encode(buf.getvalue()).decode('utf-8')
    plt.close()

    # -----------------------------
    # Sentiment Count Plot (Bar Plot)
    # -----------------------------
    custom_palette = {
        'Positive': '#90ee90',  # Light green
        'Negative': '#f08080',  # Light red
        'Neutral': '#d3d3d3'  # Light grey
    }
    plt.figure(figsize=(8, 5))
    sns.countplot(
        data=sentence_df,
        x='Sentiment',
        palette=custom_palette  # Custom color palette for the count plot
    )
    plt.title("Number of Sentences by Sentiment Distribution")
    plt.xticks(rotation=45)
    plt.grid(axis='y', linestyle='--', alpha=0.6)
    plt.tight_layout()

    buf2 = BytesIO()
    plt.savefig(buf2, format='png')
    buf2.seek(0)
    sentiment_count_plot_data = base64.b64encode(buf2.getvalue()).decode('utf-8')
    plt.close()

    # -----------------------------
    # Sentiment Distribution Plot (Histogram)
    # -----------------------------
    sentiment_data = sentence_df[sentence_df['Sentiment'] != 'Neutral']
    if not sentiment_data.empty:
        plt.figure(figsize=(8, 5))
        sns.histplot(
            data=sentiment_data,
            x='Sentiment_Score',
            hue='Sentiment',
            bins=10,  # Adjust number of bins for more or less granularity
            palette=custom_palette,  # Custom color palette for the histogram
        )
        plt.title("Sentiment Distribution of Sentences: Positive vs Negative")
        plt.xlabel("Sentiment Score")
        plt.ylabel("Frequency")
        plt.tight_layout()

        buf = BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        sentiment_plot_data = base64.b64encode(buf.getvalue()).decode('utf-8')
        plt.close()

    # Backward compatibility
    return (
        sentence_df,
        image_data,  # Returning the Readability Chart as the image_data
        sentiment_plot_data,
        sentiment_count_plot_data,
        formatted_top_positive_sentences,
        formatted_low_rank_sentences,
        formatted_top_negative_sentences
    )


@csrf_exempt
def readability_view(request):
    image_data = None
    formatted_low_rank_sentences = []
    sentences = []

    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        user_threshold = int(request.POST.get('user_threshold', 5))

        if uploaded_file:
            # Temporarily save the uploaded file
            with NamedTemporaryFile(delete=False) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            sentence_df, image_data, _, _, _, formatted_low_rank_sentences, _ = calculate_readability(
                temp_file_path, 2, 2, user_threshold)

            sentences = sentence_df.to_dict(orient='records')  # Converting DataFrame to list of dicts for use in HTML

            # Delete the temporary file after processing
            os.remove(temp_file_path)

    return render(request, 'readability.html', {
        'image_data': image_data,
        'formatted_low_rank_sentences': formatted_low_rank_sentences,
        'sentences': sentences,
    })

@csrf_exempt
def sentiment_view(request):
    sentiment_plot_data = None
    sentiment_count_plot_data = None
    formatted_top_positive_sentences = []
    formatted_top_negative_sentences = []
    sentences = []

    if request.method == 'POST':
        uploaded_file = request.FILES.get('file')
        top_positive_percentage = int(request.POST.get('top_positive', 5))
        most_negative_percentage = int(request.POST.get('most_negative', 5))

        if uploaded_file:
            # Temporarily save the uploaded file
            with NamedTemporaryFile(delete=False) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            _,_, sentiment_plot_data, sentiment_count_plot_data, formatted_top_positive_sentences, _, formatted_top_negative_sentences = calculate_readability(
                temp_file_path, top_positive_percentage, most_negative_percentage, 2)

            sentences = []  # Add logic here if you need to include sentence details as well

            # Delete the temporary file after processing
            os.remove(temp_file_path)

    return render(request, 'sentiment.html', {
        'sentiment_plot_data': sentiment_plot_data,
        'sentiment_count_plot_data': sentiment_count_plot_data,
        'formatted_top_positive_sentences': formatted_top_positive_sentences,
        'formatted_top_negative_sentences': formatted_top_negative_sentences,
        'sentences': sentences,
    })


# 🧼 Text Preprocessing Function
pattern = r'\b(?:' + '|'.join(stop_words) + r')\b'

def preprocess(text):
    text = re.sub(r'\W', ' ', str(text))      # Remove punctuation
    text = re.sub(r'\s+', ' ', text)          # Remove extra spaces
    text = re.sub(r'\d+', '', text)           # Remove numbers
    text = re.sub(r'[^a-zA-Z\s]', '', text)   # Remove special characters
    text = text.lower()                       # Convert to lowercase
    # text = re.sub(pattern, '', text, flags=re.IGNORECASE)  # Optional: stop words
    return text


# 📊 Convert df.info() to Bootstrap Table + Extra Info
def df_info_to_bootstrap(df):
    buffer = StringIO()
    df.info(buf=buffer)
    info_lines = buffer.getvalue().split('\n')

    # ✅ Extract column info for table
    table_data = []
    for line in info_lines[5:]:
        if line.strip() == '' or 'dtypes:' in line.lower():
            continue
        parts = line.split()
        if len(parts) >= 3:
            table_data.append({
                'Index': parts[0],
                'Column': parts[1],
                'Non-Null Count': parts[2],
                'Dtype': parts[-1]
            })

    # ✅ Build table
    df_info_table = pd.DataFrame(table_data).to_html(
        classes='table table-bordered table-striped',
        index=False,
        justify='left'
    )

    # ✅ Extract dtypes & memory usage
    dtypes_line = ''
    memory_line = ''
    for line in info_lines:
        if 'dtypes:' in line.lower():
            dtypes_line = line.strip()
        if 'memory usage' in line.lower():
            memory_line = line.strip()

    extra_info = f"<p><strong>{dtypes_line}</strong></p><p><strong>{memory_line}</strong></p>"
    return df_info_table, extra_info


# 📥 Main View
@csrf_exempt
def tfidf_analysis_view(request):
    df_info = None
    df_info_table = None
    df_info_extra = None
    df_sample = None
    shape_text = None
    value_counts_table = None
    histogram_image = None
    df_head_table = None
    df_sample_one_table = None
    classification_report_html = None
    conf_matrix_html = None
    pred_table_html = None
    accuracy = None
    importance_type = None
    shaped_text = None
    df_shape = None
    # Initialize these to prevent UnboundLocalError
    feature_importance_html = ""
    sorted_feature_importance_html = ""

    # ✅ Handle POST request with uploaded file
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        df = pd.read_excel(excel_file)

        # ✅ df.info() before cleaning
        if not df.empty:
            buffer = StringIO()
            df.info(buf=buffer)
            info_lines = buffer.getvalue().split('\n')

            info_data = []
            for line in info_lines[5:]:
                if line.strip() == '':
                    continue
                parts = line.split()
                if len(parts) >= 3:
                    info_data.append({
                        'Index': parts[0],
                        'Column': parts[1],
                        'Non-Null Count': parts[2],
                        'Dtype': parts[-1]
                    })

            if info_data:
                df_info_df = pd.DataFrame(info_data)
                df_info = df_info_df.to_html(
                    classes='table table-bordered table-striped',
                    index=False,
                    justify='left'
                )

        # ✅ Drop NaN rows
        df = df.dropna()

        # ✅ df.info() after cleaning
        if not df.empty:
            df_info_table, df_info_extra = df_info_to_bootstrap(df)

        # ✅ Preprocess text
        if 'Text' in df.columns:
            df['Text'] = df['Text'].apply(preprocess)
            df['word_count'] = df['Text'].apply(lambda x: len(str(x).split()))
            df['Text2'] = df['Text'].apply(preprocess)
        else:
            df['Text'] = ""
            df['Text2'] = ""

        # ✅ Samples and shape info
        if not df.empty:
            df_sample = df.sample(min(5, len(df))).to_html(
                classes='table table-bordered table-striped',
                index=False,
                justify='left'
            )

        shape_text = f"Rows: {df.shape[0]}, Columns: {df.shape[1]}"

        # ✅ Target column summary
        if 'Target' in df.columns:
            value_counts_df = df['Target'].value_counts().reset_index()
            value_counts_df.columns = ['Target', 'Count']
            value_counts_table = value_counts_df.to_html(
                classes='table table-hover table-striped',
                index=False,
                justify='left'
            )
        else:
            value_counts_table = "<p class='text-danger'>⚠️ 'Target' column not found.</p>"

        # ✅ df.head() preview
        df_head_table = df.head().to_html(
            classes="table table-bordered table-hover table-striped text-nowrap align-middle",
            index=True,
            border=0,
            justify="center"
        )

        # ✅ df.sample(1) preview
        df_sample_one_table = df.sample(1).to_html(
            classes="table table-bordered table-hover table-striped text-nowrap align-middle",
            index=True,
            border=0,
            justify="center"
        )

        # ✅ Histogram of word count
        if 'word_count' in df.columns:
            fig = px.histogram(df, x="word_count", title="Distribution of Word Counts")
            fig.update_layout(xaxis_title="Number of Words", yaxis_title="Frequency")
            img_bytes = fig.to_image(format="png")
            histogram_image = base64.b64encode(img_bytes).decode()

        # ✅ Only proceed with ML if required columns exist
        if 'Target' in df.columns and 'Text2' in df.columns:
            labels = df['Target']
            documents = df['Text2']

            vectorizer = TfidfVectorizer(ngram_range=(1, 4))
            X = vectorizer.fit_transform(documents)

            X_train, X_test, y_train, y_test = train_test_split(
                X, labels, test_size=0.25, random_state=42
            )

            model = LogisticRegression(max_iter=200)
            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)

            # ✅ Generate feature importance table
            coefficients = model.coef_[0]
            feature_names = vectorizer.get_feature_names_out()

            feature_importance = pd.DataFrame({
                'Feature': feature_names,
                'Coefficient': coefficients
            })
            feature_importance['Coefficient'] = feature_importance['Coefficient'].round(2)
            feature_importance['Absolute Coefficient'] = feature_importance['Coefficient'].abs()

            # Sort by absolute importance (highest impact first)
            feature_importance = feature_importance.sort_values(by='Absolute Coefficient', ascending=False)

            # Convert to HTML for rendering
            feature_importance_html = feature_importance.head(20).to_html(
                classes="table table-bordered table-hover table-striped text-nowrap align-middle",
                index=False,
                justify="center"
            )

            # ✅ Extended Phrase Importance Table with Row Count
            feature_importance['Target'] = np.where(
                feature_importance['Coefficient'] > 0, 'Is Target',
                np.where(feature_importance['Coefficient'] < 0, 'Not_Target', 'Other')
            )

            feature_importance['word_count'] = feature_importance['Feature'].apply(lambda x: len(x.split()))

            sorted_feature_importance = feature_importance.sort_values(
                by='Absolute Coefficient', ascending=False
            ).reset_index(drop=True)

            # Rename columns for clarity
            sorted_feature_importance = sorted_feature_importance.rename(
                columns={
                    "Feature": "Phrase",
                    "word_count": "Number of Words in Phrase"
                }
            )

            # ✅ Add a counting column (starting from 1)
            sorted_feature_importance.insert(0, "No.", range(1, len(sorted_feature_importance) + 1))

            # Convert to HTML (show top 50 rows)
            sorted_feature_importance_html = sorted_feature_importance[
                ['No.', 'Phrase', 'Coefficient', 'Target', 'Number of Words in Phrase']
            ].head(50).to_html(
                classes="table table-bordered table-hover table-striped text-nowrap align-middle",
                index=False,
                justify="center"
            )

            accuracy = accuracy_score(y_test, y_pred)
            report = classification_report(y_test, y_pred, output_dict=True)
            conf_matrix = confusion_matrix(y_test, y_pred)

            report_df = pd.DataFrame(report).transpose()
            classification_report_html = report_df.to_html(
                classes="table table-bordered table-hover table-striped text-nowrap",
                float_format="%.3f"
            )

            conf_matrix_df = pd.DataFrame(conf_matrix)
            conf_matrix_html = conf_matrix_df.to_html(
                classes="table table-bordered table-hover table-striped text-nowrap"
            )

            pred_df = pd.DataFrame({
                'Actual': y_test.values,
                'Predicted': y_pred
            }).reset_index(drop=True)

            pred_table_html = pred_df.head(20).to_html(
                classes='table table-bordered table-striped text-center',
                index=True
            )

            importance_type = str(type(model.coef_[0]))
            shaped_text = str(X.shape)
            df_shape = str(df.shape)
        else:
            accuracy = 0
            classification_report_html = "<p class='text-danger'>⚠️ Required columns ('Target', 'Text2') not found.</p>"
            conf_matrix_html = ""
            pred_table_html = ""
            importance_type = "N/A"
            shaped_text = "N/A"
            df_shape = "N/A"

    # ✅ For GET or invalid file upload, just render empty page
    return render(request, 'tfidf_analysis.html', {
        'df_info': df_info,
        'df_info_table': df_info_table,
        'df_info_extra': df_info_extra,
        'df_sample': df_sample,
        'shape_text': shape_text,
        'value_counts_table': value_counts_table,
        'histogram_image': histogram_image,
        'df_head_table': df_head_table,
        'df_sample_one_table': df_sample_one_table,
        'accuracy': round(accuracy * 100, 2) if accuracy else None,
        'classification_report_html': classification_report_html,
        'conf_matrix_html': conf_matrix_html,
        'pred_table_html': pred_table_html,
        'importance_type': importance_type,
        'shaped_text': shaped_text,
        'df_shape': df_shape,
        'feature_importance_html': feature_importance_html,
        'sorted_feature_importance_html': sorted_feature_importance_html,
    })
